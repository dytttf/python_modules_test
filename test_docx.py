#coding:utf8
import docx


#写操作
document = docx.Document()
document.add_heading("Document Title", 0)

p = document.add_paragraph('A plain paragraph having some')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True



document.save('demo.docx')
